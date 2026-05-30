# table_processor.py
import re
from docx.shared import Cm
from docx_renderer import _set_run_font, _HAS_OMML
from core.docx_reader import _format_table

try:
    from latex_to_omml import replace_latex_with_omml
except ImportError:
    def replace_latex_with_omml(para): return False


def add_markdown_table_to_doc(doc, table_lines):
    """
    处理 Markdown 表格渲染的插件函数
    """

    parsed_rows = []
    for line in table_lines:
        core_chars = set(line.strip().replace('|', '').replace(':', '').replace(' ', ''))
        if core_chars == {'-'} or not core_chars:
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(parsed_rows):
        for j, text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = text
                if cell.paragraphs and cell.paragraphs[0].runs:
                    _set_run_font(cell.paragraphs[0].runs[0],
                                  cn_font='宋体', en_font='Times New Roman',
                                  size_pt=10.5, bold=False)

                # 触发表格内公式渲染
                if _HAS_OMML and ('$' in text or '\\(' in text):
                    for p in cell.paragraphs:
                        replace_latex_with_omml(p)

    _format_table(table._element)
