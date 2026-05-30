#!/usr/bin/env python3

"""
文本 / Word → 严格排版 Word 文档转换器

支持两种输入模式：
    1. .md / .txt → 按 Markdown 标记（# / ## / ### / ```）解析并排版
    2. .docx      → 读取已有文档，自动识别段落角色后重新套用严格格式

依赖：
    pip install python-docx

用法：
    from format_conversion import convert_markdown_to_docx, reformat_docx

    convert_markdown_to_docx('input.md', 'output.docx')
    reformat_docx('input.docx', 'output.docx')

    或直接运行（自动按扩展名选模式）：
    python format_conversion.py input.md output.docx
    python format_conversion.py input.docx output.docx
"""

import os

import sys

import re

from docx import Document

from docx.shared import Pt, Cm, RGBColor

from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn

from docx.oxml import OxmlElement

try:

    from latex_to_omml import replace_latex_with_omml

    _HAS_OMML = True

except ImportError:

    _HAS_OMML = False

    def replace_latex_with_omml(para): return False

# ── 从 core 子模块导入 ───────────────────────────────────────────
from core.numbering_engine import (_ensure_list_numbering, _apply_list_numpr,
                                    _get_heading_num_id, _register_heading_numbering)
from core.latex_to_text import _latex_to_text
from core.docx_reader import (_MONOSPACE_FONTS,
                               _looks_like_code, _para_has_image, _format_table,
                               _build_numbering_maps, _container_numpr,
                               _find_numbering_lvl)
from core.ai_normalizer import _clean_pasted_text, _normalize_ai_markdown
from core.title_dictionary import TitleDictionary
from core.scene.schema import SceneConfig

# ---------------------------------------------------------------------------

# ── 从 docx_renderer 导入底层渲染函数 ──
from docx_renderer import (_set_run_font, _set_para_format, _copy_runs,
                            _strip_heading_prefix, _add_heading, _add_body,
                            _add_cover_page, _add_code_line, _HAS_OMML,
                            replace_latex_with_omml,
                            _apply_table_borders)

from dataclasses import dataclass, field
from typing import Optional, Any


@dataclass
class DocIRBlock:
    type: str
    text: str = ""
    section_type: str = "body"
    level: int = 0
    source_para: Optional[Any] = None
    source_element: Optional[Any] = None
    extra_prefix: str = ""
    meta: dict = field(default_factory=dict)


def _is_auto_numbered(para):
    """检测段落是否启用了 Word 原生自动编号"""
    if para is None or para._element is None:
        return False
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


# ── OMML 原生公式探测 ──────────────────────────────────────────
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _para_has_omml(para):
    """检测段落底层 XML 中是否包含 Word 原生公式 (OMML)"""
    if para is None or para._element is None:
        return False
    return len(para._element.findall(f'.//{{{_M_NS}}}oMath')) > 0


# ── 动态标题模式生成 ────────────────────────────────────────────

_CN_NUM_PLAIN = r'(?:[一二三四五六七八九十]{1,3})'


def get_dynamic_heading_patterns(title_dict: TitleDictionary):
    """基于 TitleDictionary 动态构建标题检测正则列表，替代静态 _HEADING_PATTERNS_READ"""
    front_pipe = title_dict.generate_front_matter_pipe()
    patterns = [
        (re.compile(r'^\s*第(\d+|[一二三四五六七八九十]+)章'), 1),
        (re.compile(r'^\s*(' + front_pipe + r')'), 1),
        (re.compile(r'^\s*(\d+(?:\.\d+){2,})[\.\、\)）]?\s*'), 3),
        (re.compile(r'^\s*(\d+\.\d+)[\.\、\)）]?\s*'), 2),
        (re.compile(r'^\s*(' + _CN_NUM_PLAIN + r')[、．]\s*'), 1),
    ]
    for item in title_dict.numbered_patterns:
        patterns.append((re.compile(item.pattern), item.level))
    return patterns


def _detect_para_role(para, text_override=None, num_maps=None, title_dict=None):
    """检测段落角色：heading(标题), code(代码), list(列表), body(正文)"""
    text = text_override if text_override is not None else (para.text if para else "")

    # 1. XML 底层自动编号解析
    if num_maps and _container_numpr(para):
        num_id, ilvl = _container_numpr(para)
        lvl_def = _find_numbering_lvl(num_maps, num_id, ilvl)
        if lvl_def:
            auto_fmt = lvl_def.get("numFmt", "")
            tpl = lvl_def.get("lvlText", "")
            if auto_fmt == "chineseCountingThousand":
                return ('heading', 1)
            if auto_fmt == "decimal" and "%" in tpl:
                level = tpl.count("%")
                if level >= 2:
                    return ('heading', level)
                return ('list',)
            return ('list',)
    elif _is_auto_numbered(para):
        return ('list',)

    # 2. 手动列表检测 (兼容 Markdown 和纯文本)
    if re.match(r'^\s*(?:\(|（)?\d+[\.\、\)）]\s+', text) or re.match(r'^\s*[-*•●○■·]\s+', text):
        return ('list',)

    # 3. 代码检测
    if para is not None and para.runs:
        run = para.runs[0]
        fn = (run.font.name or '').lower()
        size = run.font.size
        indent = para.paragraph_format.first_line_indent
        if fn in _MONOSPACE_FONTS or (size and (indent is None or indent == 0) and _looks_like_code(text)):
            return ('code',)

    # 4. 标题检测——动态模式，跟随预设
    td = title_dict or TitleDictionary(SceneConfig())
    heading_patterns = get_dynamic_heading_patterns(td)
    for pattern, level in heading_patterns:
        if pattern.match(text):
            return ('heading', level)

    # 5. 默认正文
    return ('body',)

def _apply_caption_style(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True

def _apply_reference_style(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = Cm(-0.74)
    para.paragraph_format.left_indent = Cm(0.74)

def _apply_normal_style(para):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(0.85)

def apply_style_by_role(para, role):
    style_map = {
        "caption":   _apply_caption_style,
        "reference": _apply_reference_style,
        "body":      _apply_normal_style,
    }
    handler = style_map.get(role, _apply_normal_style)
    handler(para)

def _detect_suspicious_vars(text):
    """异常变量扫描（供 GUI 面板日志使用）"""
    results = []
    lines = text.split('\n')
    for i, line in enumerate(lines, 1):
        if re.search(r'[A-Za-z]_[A-Za-z]', line) and '$' not in line:
            results.append((i, line.strip(), "疑似漏写公式符号 $...$"))
    return results

# ---------------------------------------------------------------------------
# 文本清洗与 AI 归一化管线 (DeepSeek, 豆包, ChatGPT 等)
# ---------------------------------------------------------------------------

# 核心排版主引擎
# ---------------------------------------------------------------------------

def convert_text_to_docx(text, output_file, formula_stats=None):
    """
    纯文本模式流水线：AST 抽象语法树二阶段渲染引擎
    Stage 1: markdown/block_parser 解析 → AST
    Stage 2: word_render 逐块渲染 → .docx

    返回值：严格二元组 (ChangeTracker, FormulaRuleStats)
    """
    from analyzer.change_tracker import ChangeTracker
    from core.formula_stats import FormulaRuleStats
    from md_parser.ir import BlockType
    from md_parser.block_parser import parse_markdown_text
    from md_parser.word_render import write_spans, apply_blockquote_border

    tracker = ChangeTracker()
    if formula_stats is None:
        formula_stats = FormulaRuleStats()

    from core.config_manager import get_active_scene_config, load_full_config
    scene_cfg = get_active_scene_config()
    _global_cfg = load_full_config()

    def _scfg(slot, attr, default):
        return getattr(scene_cfg.styles.get(slot, scene_cfg.styles['normal']), attr, default)

    normal_style = scene_cfg.styles["normal"]
    _en_font = normal_style.font_en
    _cn_font = normal_style.font_cn
    _size_pt = Pt(normal_style.size_pt)

    from core.heading_numbering import HeadingNumberingManager
    _heading_styles = _global_cfg.get("heading_styles", {
        "1": "cn_lower_chapter", "2": "arabic_dotted", "3": "arabic", "4": "arabic", "5": "circled"
    })
    numbering_manager = HeadingNumberingManager(_heading_styles)

    from core.ai_normalizer import hide_escapes
    text, escapes_map = hide_escapes(text)

    text = _clean_pasted_text(text)
    text = _normalize_ai_markdown(text)
    from core.formula_normalize import auto_wrap_and_normalize_context
    text = auto_wrap_and_normalize_context(text)

    # ── Stage 1: 词法与语法解析 ──
    blocks = parse_markdown_text(text)

    # 预收集脚注定义
    footnote_defs = {}
    for b in blocks:
        if b.type == BlockType.FOOTNOTE_DEF:
            footnote_defs[b.list_marker] = b.raw_text

    # ── Stage 2: AST 结构化渲染 ──
    doc = Document()
    _add_cover_page(doc, title="")

    for block_idx, block in enumerate(blocks):
        if block.type == BlockType.FOOTNOTE_DEF:
            continue

        elif block.type == BlockType.HEADING:
            prefix = numbering_manager.get_next_number(block.level)
            _add_heading(doc, block.raw_text, block.level,
                         extra_prefix=f"{prefix} " if prefix else '',
                         config=scene_cfg, global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.PARAGRAPH:
            para = doc.add_paragraph()
            from core.advanced_formatter import TypographyEngine
            TypographyEngine.apply_paragraph_style(para, normal_style)
            write_spans(para, block.spans, en_font=_en_font, cn_font=_cn_font, base_font_size=_size_pt,
                        escapes_map=escapes_map, footnote_defs=footnote_defs, doc=doc,
                        global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.BLOCKQUOTE:
            para = doc.add_paragraph()
            apply_blockquote_border(para, block.quote_level)
            _set_para_format(para, line_spacing_pt=_scfg('normal','line_spacing_pt',20), first_line_indent=Cm(0))
            write_spans(para, block.spans, en_font=_en_font, cn_font=_cn_font, base_font_size=_size_pt,
                        escapes_map=escapes_map, clear_existing=False,
                        global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.CODE_BLOCK:
            for c_line in block.code_lines:
                _add_code_line(doc, c_line, config=scene_cfg,
                               global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.LIST_ITEM:
            para = doc.add_paragraph()
            step = _scfg('normal','left_indent_cm',0.85); para.paragraph_format.left_indent = Cm(step + block.list_level * step)
            para.paragraph_format.first_line_indent = Cm(-step)
            run = para.add_run(f"{block.list_marker} ")
            _set_run_font(run, cn_font=_scfg('normal','font_cn','宋体'), en_font=_scfg('normal','font_en','Times New Roman'), size_pt=_scfg('normal','size_pt',12))
            write_spans(para, block.spans, en_font=_en_font, cn_font=_cn_font, base_font_size=_size_pt,
                        escapes_map=escapes_map, clear_existing=False,
                        global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.TASK_LIST_ITEM:
            para = doc.add_paragraph()
            step = _scfg('normal','left_indent_cm',0.85); para.paragraph_format.left_indent = Cm(step + block.list_level * step)
            para.paragraph_format.first_line_indent = Cm(-step)
            prefix = '☑ ' if block.checked else '☐ '
            run = para.add_run(prefix)
            run.font.name = 'Segoe UI Symbol'
            write_spans(para, block.spans, en_font=_en_font, cn_font=_cn_font, base_font_size=_size_pt,
                        escapes_map=escapes_map, clear_existing=False,
                        global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.TABLE:
            from table_processor import add_markdown_table_to_doc
            header_line = '| ' + ' | '.join(block.table_headers) + ' |'
            sep_line = '|' + '|'.join(
                ':' + '-' * 3 + ':' if a == 'center'
                else '-' * 4 + ':' if a == 'right'
                else '-' * 4
                for a in block.table_alignments) + '|'
            row_lines = ['| ' + ' | '.join(r) + ' |' for r in block.table_rows]
            table_lines = [header_line, sep_line] + row_lines
            add_markdown_table_to_doc(doc, table_lines)
            if doc.tables:
                from core.docx_reader import _format_table
                from docx_renderer import _apply_table_borders
                _format_table(doc.tables[-1])
                _apply_table_borders(doc.tables[-1], config=scene_cfg)

        elif block.type == BlockType.LATEX_BLOCK:
            from docx_renderer import _add_equation_table
            _add_equation_table(doc, block.raw_text, config=scene_cfg,
                                global_stats=formula_stats, para_idx=block_idx)

        elif block.type == BlockType.HORIZONTAL_RULE:
            para = doc.add_paragraph()
            _set_para_format(para, line_spacing_pt=_scfg('normal','line_spacing_pt',20), first_line_indent=Cm(0))
            run = para.add_run('─' * 40)
            _set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=10)

        elif block.type == BlockType.BLANK:
            pass  # 跳过空行——段落间距由预设控制，不靠空段落

        # 未知类型 → 默认正文兜底
        else:
            para = doc.add_paragraph()
            _set_para_format(para, line_spacing_pt=_scfg('normal','line_spacing_pt',20), first_line_indent=Cm(_scfg('normal','first_line_indent_cm',0.85)))
            if block.raw_text:
                para.add_run(block.raw_text)

    _apply_page_setup(doc, scene_cfg)
    doc.save(output_file)
    return tracker, formula_stats

def convert_markdown_to_docx(input_file, output_file):
    """处理 .md / .txt 文件的入口"""
    from core.formula_stats import FormulaRuleStats
    formula_stats = FormulaRuleStats()
    with open(input_file, 'r', encoding='utf-8') as f:
        text = f.read()
    return convert_text_to_docx(text, output_file, formula_stats=formula_stats)

def _apply_page_setup(dst, scene_cfg):
    """将 SceneConfig 中的页面设置（纸张/页边距/分栏）写入输出文档的节属性"""
    from docx.shared import Cm as _Cm
    page = scene_cfg.page_setup
    body = dst._body._element
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        return

    _PAPER_TWIPS = {"A4": (11906, 16838), "Letter": (12240, 15840),
                    "US Letter": (12240, 15840)}

    def _ensure(parent, tag, after=None):
        el = parent.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            if after is not None:
                after_el = parent.find(qn(after))
                if after_el is not None:
                    idx = list(parent).index(after_el)
                    parent.insert(idx + 1, el)
                    return el
            parent.append(el)
        return el

    # 纸张大小
    sizes = _PAPER_TWIPS.get(page.paper_size)
    if sizes:
        pgSz = _ensure(sectPr, 'w:pgSz')
        pgSz.set(qn('w:w'), str(sizes[0]))
        pgSz.set(qn('w:h'), str(sizes[1]))

    # 页边距
    m = page.margin
    pgMar = _ensure(sectPr, 'w:pgMar', after='w:pgSz')
    pgMar.set(qn('w:top'), str(int(m.top_cm * 567)))
    pgMar.set(qn('w:bottom'), str(int(m.bottom_cm * 567)))
    pgMar.set(qn('w:left'), str(int(m.left_cm * 567)))
    pgMar.set(qn('w:right'), str(int(m.right_cm * 567)))

    # 分栏
    cols_cfg = page.columns
    if cols_cfg.count > 1:
        cols_el = _ensure(sectPr, 'w:cols', after='w:pgMar')
        cols_el.set(qn('w:num'), str(cols_cfg.count))
        if cols_cfg.space_cm > 0:
            cols_el.set(qn('w:space'), str(int(cols_cfg.space_cm * 567)))
        if cols_cfg.equal_width:
            cols_el.set(qn('w:equalWidth'), '1')
        if cols_cfg.separator:
            cols_el.set(qn('w:sep'), '1')


def reformat_docx(input_file, output_file, config=None):
    """
    重排版已有 Word 文档的入口（三层架构版）

    Layer 1: 静态封面/目录注入
    Layer 2: 模板样式 DNA 克隆
    Layer 3: 底层硬编码兜底
    """
    import copy
    from analyzer.change_tracker import ChangeTracker
    from core.formula_stats import FormulaRuleStats
    tracker = ChangeTracker()
    formula_stats = FormulaRuleStats()

    if config is None:
        from template_config import TemplateConfig
        config = TemplateConfig()

    # ── Layer 0: 免疫清洗 ──
    if input_file.lower().endswith('.docx'):
        from sanitize import sanitize_docx
        input_file = sanitize_docx(input_file)

    src = Document(input_file)

    # ── MathType OLE 预处理：扫描并洗白 MathType 公式为 $LaTeX$ ──
    from core.ole_washer import wash_mathtype_oles
    cleaned = wash_mathtype_oles(src)
    if cleaned > 0:
        tracker.record(
            rule_name="MathType OLE 清洗", target="全文档", section="global",
            change_type="convert", before="MathType OLE 二进制对象",
            after=f"提取并洗白 {cleaned} 个公式",
            paragraph_index=-1, success=True
        )

    # ── Layer 2 / Layer 3: 完美底座构建 ──
    if config.layer_mode == 2:
        from style_shell import create_empty_shell, create_hybrid_shell
        if config.cover_path:
            dst = create_hybrid_shell(config.cover_path, config.template_path)
        else:
            dst = create_empty_shell(config.template_path)
    else:
        if config.cover_path:
            # 直接把封面文档作为底座打开，图片/样式 100% 保留
            dst = Document(config.cover_path)
            dst.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            dst = Document()
            _add_cover_page(dst, title="")

    # ── 文档结构分析 ──
    from analyzer.doc_tree import DocTree
    doc_tree = DocTree()
    doc_tree.build(src)

    num_maps = _build_numbering_maps(src)

    from core.config_manager import get_active_scene_config
    scene_cfg = get_active_scene_config()

    title_dictionary = TitleDictionary(scene_cfg)

    # O(1) 段落索引映射（替代 O(N²) 线性查找）
    para_to_idx = {p._element: i for i, p in enumerate(src.paragraphs)}

        # ── Stage 1: Parser ──
    raw_ir_stream = []
    for child in src._body._body:
        if child.tag.endswith('p'):
            para_idx = para_to_idx.get(child)
            if para_idx is None:
                continue
            para = src.paragraphs[para_idx]
            text = para.text.strip()
            section_type = doc_tree.get_section_for_paragraph(para_idx)

            if config.cover_path and section_type in ('cover','toc','abstract_cn','abstract_en','pre_body'):
                continue

            if _para_has_image(para):
                new_p = copy.deepcopy(child)
                pPr = new_p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr'); new_p.insert(0, pPr)
                old_sp = pPr.find(qn('w:spacing'))
                if old_sp is not None: pPr.remove(old_sp)
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:line'),'240'); sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
                raw_ir_stream.append(DocIRBlock(type='image', section_type=section_type, source_element=new_p))
                continue

            if _para_has_omml(para):
                new_p = copy.deepcopy(child)
                raw_ir_stream.append(DocIRBlock(type='omml', section_type=section_type, source_element=new_p))
                continue

            if not text:
                continue

            if _container_numpr(para):
                text = re.sub(
                    r'^\s*(?:[一二三四五六七八九十]+[、．]|\d+(?:\.\d+)+[\.\、\)）]?\s+|\d+[\.\、\)）]\s+)',
                    '', text)

            role = _detect_para_role(para, text_override=text, num_maps=num_maps, title_dict=title_dictionary)
            rt = role[0]; level = role[1] if len(role) > 1 else 0

            if rt == 'body':
                refined = doc_tree.get_refined_role(para_idx, text)
                if refined in ('caption','reference'): rt = refined

            lvl = 0
            if rt == 'list':
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    nPr = pPr.find(qn('w:numPr'))
                    if nPr is not None:
                        il = nPr.find(qn('w:ilvl'))
                        if il is not None: lvl = int(il.get(qn('w:val'),0))

            meta = {}
            num_info = _container_numpr(para)
            if num_info: meta['num_info'] = num_info

            raw_ir_stream.append(DocIRBlock(
                type=rt, text=text, section_type=section_type,
                level=level or lvl, source_para=para, meta=meta))

        elif child.tag.endswith('tbl'):
            raw_ir_stream.append(DocIRBlock(type='table', section_type='body', source_element=copy.deepcopy(child)))

    # ── Stage 2: IR Optimizer ──
    optimized_ir_stream = []
    for block in raw_ir_stream:
        if block.type == 'body' and block.text.strip() == '$$':
            continue

        if block.type == 'reference' or block.section_type == 'references':
            parts = re.split(r'\s+(?=\[\d{1,3}\])', block.text)
            if len(parts) > 1:
                for part in parts:
                    part = part.strip()
                    if part:
                        optimized_ir_stream.append(DocIRBlock(type='reference', text=part, section_type='references'))
                continue
        optimized_ir_stream.append(block)

    # ── Stage 3: Renderer ──
    from lxml import etree as _lx

    def _cfg(slot, attr, default):
        return getattr(scene_cfg.styles.get(slot, scene_cfg.styles['normal']), attr, default)

    for block in optimized_ir_stream:
        if block.type in ('image', 'omml'):
            sectPr = dst._body._element.find(qn('w:sectPr'))
            if sectPr is not None:
                sectPr.addprevious(block.source_element)
            else:
                dst._body._element.append(block.source_element)

        elif block.type == 'table':
            new_tbl = block.source_element
            sectPr = dst._body._element.find(qn('w:sectPr'))
            if sectPr is not None: sectPr.addprevious(new_tbl)
            else: dst._body._element.append(new_tbl)
            _format_table(new_tbl)
            _apply_table_borders(new_tbl, config=scene_cfg)
            for cell in new_tbl.iter(qn('w:tc')):
                for p in cell.iter(qn('w:p')):
                    tp = OxmlElement('w:p')
                    for c in list(p): tp.append(copy.deepcopy(c))
                    td = Document(); td._body._element.append(tp)
                    if td.paragraphs: replace_latex_with_omml(td.paragraphs[0], global_stats=formula_stats)

        elif block.type == 'heading':
            hp = _add_heading(dst, block.text, block.level, source_para=block.source_para,
                              config=config, global_stats=formula_stats, para_idx=-1)
            if 'num_info' in block.meta and block.meta['num_info'] and num_maps:
                num_id, ilvl = block.meta['num_info']
                lvl_def = _find_numbering_lvl(num_maps, num_id, ilvl)
                if lvl_def:
                    pPr = hp._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = _lx.SubElement(hp._element, qn('w:pPr'))
                        hp._element.insert(0, pPr)
                    numPr = _lx.SubElement(pPr, qn('w:numPr'))
                    nid = _lx.SubElement(numPr, qn('w:numId')); nid.set(qn('w:val'), str(num_id))
                    il = _lx.SubElement(numPr, qn('w:ilvl')); il.set(qn('w:val'), str(ilvl))

        elif block.type == 'code':
            _add_code_line(dst, block.text, source_para=block.source_para, global_stats=formula_stats, para_idx=-1)

        elif block.type == 'list':
            step = _cfg('normal', 'left_indent_cm', 0.85)
            tl = Cm(step + block.level * step)
            _add_body(dst, block.text, left_indent=tl, first_line_indent=Cm(-step),
                      source_para=block.source_para, config=config, global_stats=formula_stats, para_idx=-1)

        elif block.type == 'caption':
            pn = _add_body(dst, block.text, source_para=block.source_para, config=config, global_stats=formula_stats, para_idx=-1)
            _apply_caption_style(pn)

        elif block.type == 'reference':
            ref_left = _cfg('references_body', 'left_indent_cm', 0.74)
            ref_hang = _cfg('references_body', 'hanging_indent_cm', 0.74)
            _add_body(dst, block.text, source_para=block.source_para, config=config,
                      left_indent=Cm(ref_left), first_line_indent=Cm(-ref_hang),
                      section_type='references', global_stats=formula_stats, para_idx=-1)

        else:
            clean_text = block.text.strip().strip('$').strip()
            if (block.type == 'body'
                    and re.search(r'(?:[\(\（][0-9\-–\.]+[\)\）]|\\tag\s*\{?[0-9\-–\.]+\}?)\s*$', clean_text)
                    and re.search(r'(\$|\\|[a-zA-Z]=)', block.text)):
                from docx_renderer import _add_equation_table
                _add_equation_table(dst, block.text, config=scene_cfg,
                                    global_stats=formula_stats, para_idx=-1)
                continue

            cl = None; cf = None
            if block.section_type == 'references':
                cl = Cm(_cfg('references_body', 'left_indent_cm', 0.85))
                cf = Cm(-_cfg('references_body', 'hanging_indent_cm', 0.85))
            elif block.section_type in ('toc','cover'): cf = Cm(0)
            _add_body(dst, block.text, source_para=block.source_para, config=config,
                      left_indent=cl, first_line_indent=cf,
                      section_type=block.section_type, global_stats=formula_stats, para_idx=-1)

# ── 公式健康度审计报告 ──
    if formula_stats.matched > 0:
        low_conf_occs = [o for o in formula_stats.occurrences if o.confidence < 0.85]
        if not low_conf_occs:
            summary = f"完美处理 {formula_stats.matched} 个公式 (全高置信度)"
            tracker.record(
                rule_name="公式引擎", target="全文档", section="formula",
                change_type="convert", before=f"发现 {formula_stats.matched} 个 LaTeX 公式",
                after=summary, paragraph_index=-1, success=True
            )
        else:
            warnings_info = "; ".join([
                f"[{occ.original_text[:15]}...] -> {','.join(occ.warnings)}"
                for occ in low_conf_occs[:6]
            ])
            tracker.record(
                rule_name="公式引擎 (需复核)", target="全文档", section="formula",
                change_type="warning", before=f"存在 {len(low_conf_occs)} 个潜在语法错误",
                after=warnings_info, paragraph_index=-1, success=False,
                failure_reason="建议人工检查标红的公式"
            )

    _apply_page_setup(dst, scene_cfg)
    dst.save(output_file)

    # ── 保存后轻量级 Validation 扫描 ──
    try:
        val_doc = Document(output_file)
        for i, p in enumerate(val_doc.paragraphs):
            text = p.text
            if '$' in text or '\\frac' in text or '\\_' in text:
                tracker.record(
                    rule_name="Formula Conversion", target=f"第 {i+1} 段", section="global",
                    change_type="format", before=text[:30]+"...", after="转换失败",
                    paragraph_index=i, success=False,
                    failure_reason="存在疑似未成功转换的 LaTeX 公式标记，请检查语法。"
                )
            if '\t' in text:
                tracker.record(
                    rule_name="Spacing Validation", target=f"第 {i+1} 段", section="global",
                    change_type="text", before="存在 Tab 键", after="跳过",
                    paragraph_index=i, success=False,
                    failure_reason="段落中包含 Tab 键，可能导致对齐错乱，建议删除。"
                )
    except Exception:
        pass

    return tracker, formula_stats

# ---------------------------------------------------------------------------
# CLI 命令行运行入口
# ---------------------------------------------------------------------------
if __name__ == '__main__':
    print("=========================================")
    print(" FormatX - 全能 AI 格式排版引擎已启动")
    print("=========================================")
    if len(sys.argv) >= 3:
        input_file, output_file = sys.argv[1], sys.argv[2]
        ext = os.path.splitext(input_file)[1].lower()
        if ext == '.docx':
            tracker, formula_stats = reformat_docx(input_file, output_file)
        else:
            tracker, formula_stats = convert_markdown_to_docx(input_file, output_file)
        print(f"\n转换成功，文件已保存至: {output_file}")
    elif len(sys.argv) == 2:
        input_file = sys.argv[1]
        ext = os.path.splitext(input_file)[1].lower()
        output_file = input_file.replace(ext, '_formatted.docx')
        if ext == '.docx':
            tracker, formula_stats = reformat_docx(input_file, output_file)
        else:
            tracker, formula_stats = convert_markdown_to_docx(input_file, output_file)
        print(f"\n转换成功，文件已保存至: {output_file}")
    else:
        print("\n【使用方法】")
        print("1. 拖拽文件到 exe 上直接运行。")
        print("2. 命令行: python format_conversion.py <输入文件> <输出文件>")
        sys.exit(0)

    from core.report.collector import collect_report
    from core.report.markdown_report import generate_markdown_report
    report_data = collect_report(tracker, formula_stats, input_file=input_file)
    report_md = output_file.replace('.docx', '_报告.md')
    generate_markdown_report(report_data, report_md)
    print(f"审计报告: {report_md}")
