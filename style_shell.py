"""
样式空壳提取器与封面拼接器

Layer 2: 提取模板文档的样式 DNA，清空内容返回空壳
Layer 1: 将封面文档的 XML 节点安全深拷贝到目标文档前
"""
import copy
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn


def create_empty_shell(template_path):
    """【Layer 2】提取模板的样式灵魂，清空肉体，返回空壳文档"""
    doc = Document(template_path)
    body = doc._body._element
    sectPr = body.find(qn('w:sectPr'))

    for child in list(body):
        if child != sectPr:
            body.remove(child)
    return doc


def inject_cover(doc, cover_path):
    """【Layer 1】物理拼接封面/目录，使用安全的深拷贝法"""
    if not cover_path:
        return

    cover_doc = Document(cover_path)
    dst_body = doc._body._element
    dst_sectPr = dst_body.find(qn('w:sectPr'))

    for child in cover_doc._body._element:
        if child.tag.endswith('sectPr'):
            continue
        new_child = copy.deepcopy(child)
        if dst_sectPr is not None:
            dst_sectPr.addprevious(new_child)
        else:
            dst_body.append(new_child)

    # 封面后打入分页符，保证正文另起一页
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def create_hybrid_shell(cover_path, template_path):
    """
    【双轨融合】以封面为肉体（保留校徽/下划线），以模板为灵魂（注入正文排版）
    """
    # 1. 打开封面作为绝对物理底座（保住 media 图片和 .rels 关系）
    dst = Document(cover_path)

    # 2. 打开正文模板，准备抽取灵魂
    tpl = Document(template_path)

    # 3. XML 级 styles 移植：模板样式覆盖封面样式
    dst_styles = dst.styles.element
    tpl_styles = copy.deepcopy(tpl.styles.element)
    dst_styles.getparent().replace(dst_styles, tpl_styles)
    # 强制刷新 python-docx 内部缓存
    dst.styles._element = tpl_styles

    # 4. 移植多级列表编号基因 (numbering.xml)
    try:
        try:
            tpl_num_part = tpl.part.numbering_part
        except Exception:
            tpl_num_part = None

        if tpl_num_part:
            tpl_num = copy.deepcopy(tpl_num_part._element)
            try:
                dst_num_part = dst.part.numbering_part
            except Exception:
                dst_num_part = None

            if dst_num_part:
                dst_num = dst_num_part._element
                dst_num.getparent().replace(dst_num, tpl_num)
    except Exception:
        pass

    # 5. 在封面末尾打入分页符
    dst.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    return dst
