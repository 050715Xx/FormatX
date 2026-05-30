"""
FormatX 底层 Word 渲染器
字体设置 / 段落格式 / Run复制 / 标题 / 正文 / 封面 / 代码行
"""
import re
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

from core.latex_to_text import _latex_to_text
from core.numbering_engine import _get_heading_num_id

try:
    from latex_to_omml import replace_latex_with_omml
    _HAS_OMML = True
except ImportError:
    _HAS_OMML = False
    def replace_latex_with_omml(para): return False


# 底层格式工具

def _set_xml_border_item(parent_el, edge, sz=0, val='single', color='auto', space=0):
    """操纵底层 XML 边框的原子构建器 (sz 的单位是 1/8 磅)"""
    border = OxmlElement(f'w:{edge}')
    border.set(qn('w:val'), val)
    if val != 'none':
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), color)
    parent_el.append(border)


def _apply_table_borders(table, config=None):
    """学术级表格边框专控引擎——作为独立装饰层叠加，不破坏单元格内边距与字体"""
    if config is None:
        from core.scene.schema import SceneConfig
        config = SceneConfig()

    mode = getattr(config, "normal_table_border_mode", "three_line")
    top_bottom = getattr(config, "three_line_header_width_pt", 1.0)
    mid = getattr(config, "three_line_bottom_width_pt", 0.5)
    full = getattr(config, "table_border_width_pt", 0.5)

    tb_sz = int(top_bottom * 8)
    mid_sz = int(mid * 8)
    full_sz = int(full * 8)

    # 兼容 Table 对象和原始 CT_Tbl XML 元素
    tbl_el = table._tbl if hasattr(table, '_tbl') else table
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)

    # 拔掉 TableGrid 等默认样式，确保直接格式化边框优先生效
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)

    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    new_borders = OxmlElement('w:tblBorders')

    if mode == "three_line":
        # 无差别清空所有单元格旧边框，防止源文档残留优先级覆盖
        for row in tbl_el.findall(qn('w:tr')):
            for tc in row.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    old_tc_b = tcPr.find(qn('w:tcBorders'))
                    if old_tc_b is not None:
                        tcPr.remove(old_tc_b)
        _set_xml_border_item(new_borders, 'top', sz=tb_sz, val='single')
        _set_xml_border_item(new_borders, 'bottom', sz=tb_sz, val='single')
        for edge in ('left', 'right', 'insideH', 'insideV'):
            _set_xml_border_item(new_borders, edge, val='none')
        tblPr.append(new_borders)
        # 表头行底部细线
        rows_el = tbl_el.findall(qn('w:tr'))
        if rows_el:
            for tc in rows_el[0].findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                old_tc = tcPr.find(qn('w:tcBorders'))
                if old_tc is not None:
                    tcPr.remove(old_tc)
                tcBorders = OxmlElement('w:tcBorders')
                _set_xml_border_item(tcBorders, 'bottom', sz=mid_sz, val='single')
                tcPr.append(tcBorders)
    elif mode == "full_grid":
        for row in tbl_el.findall(qn('w:tr')):
            for tc in row.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    old_tc_b = tcPr.find(qn('w:tcBorders'))
                    if old_tc_b is not None:
                        tcPr.remove(old_tc_b)
        for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
            _set_xml_border_item(new_borders, edge, sz=full_sz, val='single')
        tblPr.append(new_borders)


def _set_no_proof(run):
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn('w:noProof')) is None:
        rpr.append(OxmlElement('w:noProof'))


# ---------------------------------------------------------------------------

def _set_run_font(run, cn_font='宋体', en_font='Times New Roman',

                  size_pt=12, bold=False):

    """

    为 run 同时设置中文字体（w:eastAsia）和西文字体（w:ascii / w:hAnsi）。

    这是整个排版的关键——python-docx 的 run.font.name 只能设西文，

    中文必须通过 XML 层的 w:eastAsia 属性单独指定。

    """

    run.font.size = Pt(size_pt)

    run.bold = bold

    run.font.color.rgb = RGBColor(0, 0, 0)  # 强制黑色，屏蔽 Heading 样式默认蓝色

    # 西文字体（影响 ASCII 数字、英文、符号）

    run.font.name = en_font

    # 东亚字体（影响中文、日文、韩文）

    rPr = run._r.get_or_add_rPr()

    rFonts = rPr.find(qn('w:rFonts'))

    if rFonts is None:

        rFonts = OxmlElement('w:rFonts')

        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), cn_font)

def _set_para_format(para, line_spacing_pt=20, first_line_indent=None):

    """

    设置段落行距（固定值）和首行缩进。

    参数

    ----

    line_spacing_pt : float

        固定行距，单位磅。Pt(20) 即为「固定值 20 磅」。

    first_line_indent : Cm | Pt | None

        首行缩进距离。传 Cm(0) 即明确取消缩进（代码块）；

        传 Cm(0.85) 即正文两字符缩进；None 表示不修改。

    """

    pf = para.paragraph_format

    pf.line_spacing = Pt(line_spacing_pt)

    # 去掉段前段后间距，保证严格 20 磅行距栅格

    pf.space_before = Pt(0)

    pf.space_after = Pt(0)

    if first_line_indent is not None:

        pf.first_line_indent = first_line_indent

# ---------------------------------------------------------------------------

# 段落构建辅助

# ---------------------------------------------------------------------------

def _copy_runs(para, source_para, cn_font, en_font, size_pt, default_bold,
               config=None):

    """

    从 source_para 逐 run 复制到 para，保留原有的 bold/italic/underline，
    仅替换字体名和字号。Layer 2 模式下跳过硬编码字体，让模板 DNA 接管。

    """

    if source_para is None:

        return None  # 调用方自行处理

    # 兜底：获取源段落样式本身的加粗倾向
    p_bold = source_para.style.font.bold if (source_para.style and source_para.style.font) else False

    # 段落级数学公式嗅探：只要含 $ 或 \( 或 \[，全段原样保留
    has_math = bool(re.search(r'(\$|\\\[|\\\(|\\[A-Za-z]+(?![a-zA-Z]))', source_para.text))

    for src_run in source_para.runs:

        # OMML 模式：只要段落里有公式，全段保持原样，绝不走纯文本降级
        if _HAS_OMML and has_math:
            run_text = src_run.text
        elif _HAS_OMML and '$' in src_run.text:

            run_text = src_run.text

        else:

            run_text = _latex_to_text(src_run.text)

        run = para.add_run(run_text)

        # 四级回退：run 自身 → 字符样式 → 段落样式 → 默认值
        r_style_bold = src_run.style.font.bold if (src_run.style and src_run.style.font) else None

        if src_run.bold is not None:
            actual_bold = src_run.bold
        elif r_style_bold is not None:
            actual_bold = r_style_bold
        else:
            actual_bold = p_bold if p_bold is not None else default_bold

        run.bold = actual_bold

        if src_run.italic is not None:

            run.italic = src_run.italic

        if src_run.underline is not None:

            run.underline = src_run.underline

        # Layer 2 模式下跳过硬编码字体，让模板 DNA 接管
        if config is None or config.layer_mode != 2:
            _set_run_font(run, cn_font=cn_font, en_font=en_font,
                          size_pt=size_pt, bold=run.bold)

    return para

def _strip_heading_prefix(text, level):

    """【所见即所得】：绝对不切除任何前缀，原样保留用户复制的内容。"""

    return text.strip()

def _add_heading(doc, text, level, alignment=None, source_para=None,

                 extra_prefix='', config=None, global_stats=None, para_idx=-1):

    """

    添加一级 / 二级 / 三级标题。

    - Layer 2（模板模式）：赋 Heading 样式，由模板 DNA 完全接管
    - Layer 3（兜底模式）：硬编码宋体/TNR/固定行距

    """

    para = doc.add_paragraph()

    from core.constants import STYLE_MAP
    heading_style = STYLE_MAP.get(f"heading{level}")

    is_layer2 = (config is not None and getattr(config, 'layer_mode', None) == 2)
    h_style = config.styles.get(f"heading{level}") if (config and hasattr(config, 'styles')) else None
    _h_cn = h_style.font_cn if h_style else '宋体'
    _h_en = h_style.font_en if h_style else 'Times New Roman'
    _h_sz = h_style.size_pt if h_style else {1: 16, 2: 14, 3: 12}.get(level, 12)

    # ── Layer 2：仅赋样式，不做段落格式硬编码 ──
    if is_layer2:
        if heading_style:
            try:
                para.style = doc.styles[heading_style]
                _ = para.style.style_id
            except KeyError:
                pass
    else:
        _set_para_format(para, line_spacing_pt=20, first_line_indent=Cm(0))

        if level in (1, 2):
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(10)

            if alignment is not None:
                para.paragraph_format.alignment = alignment
            elif re.match(r'^第(\d+|[一二三四五六七八九十]+)章', text.strip()) or \
                 re.match(r'^(摘要|Abstract|ABSTRACT|目录|参考文献|致谢|附录)', text.strip()):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.page_break_before = True
            else:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── 以下逻辑 Layer 2 / Layer 3 共用 ──
    title_text = re.sub(r'\s+', '', text.strip()) if re.search(r'[一-鿿]', text) else text.strip()
    title_text = re.sub(r'\*\*', '', title_text)  # 标题全段加粗，** 符号多余

    use_native = (heading_style is not None and extra_prefix == '')

    if use_native and not is_layer2:
        try:
            para.style = doc.styles[heading_style]
            _ = para.style.style_id
            _get_heading_num_id(doc)
        except KeyError:
            pass
        title_text = _strip_heading_prefix(title_text, level)

    if extra_prefix:
        pre_run = para.add_run(extra_prefix)
        if not is_layer2:
            _set_run_font(pre_run, cn_font=_h_cn, en_font=_h_en,
                          size_pt=_h_sz, bold=True)

    is_cleaned_native = use_native and title_text != text.strip()

    if source_para is not None and not is_cleaned_native:
        _copy_runs(para, source_para, _h_cn, _h_en, _h_sz, True, config=config)
    else:
        run = para.add_run(title_text)
        if not is_layer2:
            _set_run_font(run, cn_font=_h_cn, en_font=_h_en,
                          size_pt=_h_sz, bold=True)

    if _HAS_OMML:
        replace_latex_with_omml(para, global_stats=global_stats, para_idx=para_idx)
    return para

def _add_body(doc, text, alignment=None, left_indent=None,

              first_line_indent=None, source_para=None, extra_prefix='',

              style_name=None, config=None, section_type='body', global_stats=None,
              para_idx=-1):

    """

    添加正文段落。Layer 2 仅挂载 Normal 样式，不硬编码字体格式。
    section_type: 分区类型 (body / references / toc / cover 等)
    """

    first = first_line_indent if first_line_indent is not None else Cm(0.85)

    para = doc.add_paragraph()

    is_layer2 = (config is not None and getattr(config, 'layer_mode', None) == 2)

    # ── Layer 2：仅挂载样式 ──
    if is_layer2:
        target_style = 'Normal'
        if section_type == 'references':
            for s_name in ['参考文献', 'Bibliography', 'List Paragraph']:
                if s_name in doc.styles:
                    target_style = s_name
                    break
        try:
            para.style = doc.styles[target_style]
            _ = para.style.style_id
        except KeyError:
            pass

        if section_type == 'references' and target_style == 'Normal':
            para.paragraph_format.left_indent = Cm(0.85)
            para.paragraph_format.first_line_indent = Cm(-0.85)
    else:
        # ── Layer 3：硬编码段落格式 ──
        _set_para_format(para, line_spacing_pt=20, first_line_indent=first)

        if alignment is not None:
            para.paragraph_format.alignment = alignment
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if left_indent is not None:
            para.paragraph_format.left_indent = left_indent

    # ── 以下逻辑 Layer 2 / Layer 3 共用 ──
    # 保留原列表样式
    if style_name and 'list' in style_name.lower():
        try:
            para.style = doc.styles[style_name]
            _ = para.style.style_id
        except KeyError:
            pass

    # 自动编号前缀（如有）
    if extra_prefix:
        pre_run = para.add_run(extra_prefix)
        if not is_layer2:
            _set_run_font(pre_run, cn_font='宋体', en_font='Times New Roman',
                          size_pt=12, bold=True)

    if source_para is not None:

        # 联合判断：run 自身 + 段落样式继承
        has_runs = len(source_para.runs) > 0
        p_style_bold = source_para.style.font.bold if (source_para.style and source_para.style.font) else False

        def _is_run_bold(r):
            r_style_bold = r.style.font.bold if (r.style and r.style.font) else None
            if r.bold is not None: return r.bold
            if r_style_bold is not None: return r_style_bold
            if p_style_bold is not None: return p_style_bold
            return False

        all_runs_bold = has_runs and all(_is_run_bold(r) for r in source_para.runs)

        if all_runs_bold:
            full_text = source_para.text
            # 冒号雷达：整段加粗时，若存在冒号则只保留冒号前的标题加粗
            if '：' in full_text or ':' in full_text:
                delim = '：' if '：' in full_text else ':'
                parts = full_text.split(delim, 1)

                r1 = para.add_run(parts[0] + delim)
                if not is_layer2:
                    _set_run_font(r1, cn_font='宋体', en_font='Times New Roman', size_pt=12, bold=True)
                else:
                    r1.bold = True

                r2 = para.add_run(parts[1].strip())
                if not is_layer2:
                    _set_run_font(r2, cn_font='宋体', en_font='Times New Roman', size_pt=12, bold=False)
            else:
                # 无冒号全段加粗：整体降级，防止伪装成标题
                _copy_runs(para, source_para, '宋体', 'Times New Roman', 12, False, config=config)
                for r in para.runs:
                    r.bold = False
        else:
            # 局部加粗：完美继承原文档的局部格式
            _copy_runs(para, source_para, '宋体', 'Times New Roman', 12, False, config=config)

    else:

        raw_text = text.strip()
        parts = re.split(r'\*\*(.*?)\*\*', raw_text)

        for i, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            is_bold = (i % 2 != 0)

            if not is_layer2:
                _set_run_font(run, cn_font='宋体', en_font='Times New Roman',
                              size_pt=12, bold=is_bold)
            else:
                run.bold = is_bold

    # 将段落中的 $...$ LaTeX 转为 OMML 公式
    if _HAS_OMML:
        replace_latex_with_omml(para, global_stats=global_stats, para_idx=para_idx)
    return para

def _add_cover_page(doc, title=''):

    """

    添加封面预留页。

    - 如果提供了 title，则按封面标题格式渲染：

      黑体，初号（42 pt），加粗，居中

    - 如果 title 为空字符串，则留空白封面

    - 末尾插入分页符，后续内容从第 2 页开始

    """

    if doc.paragraphs:

        para = doc.paragraphs[0]

    else:

        para = doc.add_paragraph()

    if title:

        # 封面标题：黑体 初号（42 pt）加粗 居中

        run = para.add_run(title.strip())

        _set_run_font(run, cn_font='黑体', en_font='黑体',

                      size_pt=42, bold=True)

        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_para_format(para, line_spacing_pt=20, first_line_indent=Cm(0))

    para.add_run().add_break(WD_BREAK.PAGE)

def _add_code_line(doc, line_text, source_para=None, config=None,
                   global_stats=None, para_idx=-1):

    """

    添加代码块中的一行。

    - Times New Roman 5号（11 pt）

    - 不缩进，保留原始前导空格（维持缩进层级）

    - 固定行距 20 磅

    source_para 不为 None 时逐 run 复制保留原加粗/斜体。

    """

    cb_style = config.styles.get("code_block") if (config and hasattr(config, 'styles')) else None
    _cb_cn = cb_style.font_cn if (cb_style and cb_style.font_cn) else '宋体'
    _cb_en = cb_style.font_en if cb_style else 'Times New Roman'
    _cb_sz = cb_style.size_pt if cb_style else 11

    para = doc.add_paragraph()

    _set_para_format(para, line_spacing_pt=20, first_line_indent=Cm(0))

    if source_para is not None:

        _copy_runs(para, source_para, _cb_cn, _cb_en, _cb_sz, False)
        for run_obj in para.runs:
            _set_no_proof(run_obj)

    else:

        run = para.add_run(line_text)

        _set_run_font(run, cn_font=_cb_cn, en_font=_cb_en,

                      size_pt=_cb_sz, bold=False)
        _set_no_proof(run)

    # 将段落中的 $...$ LaTeX 转为 OMML 公式

    if _HAS_OMML:

        replace_latex_with_omml(para, global_stats=global_stats, para_idx=para_idx)

# ---------------------------------------------------------------------------

# .docx 输入：段落角色检测

# ---------------------------------------------------------------------------

# 常见等宽字体（用于识别代码块）

# 段落角色检测与异常扫描


# ── 公式编号表格夹具 ──────────────────────────────────────────────

def _add_equation_table(doc, text, config=None, global_stats=None, para_idx=-1):
    """构建一个完全隐形的 1×3 表格，公式居中、编号右对齐。无编号则降级为普通正文。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    clean_text = text.strip().strip('$').strip()
    match = re.search(
        r'^(.*?)\s*(?:[\(\（]([0-9\-–\.]+)[\)\）]|\\tag\s*\{?([0-9\-–\.]+)\}?)\s*$',
        clean_text, flags=re.DOTALL)
    if not match:
        return _add_body(doc, text, global_stats=global_stats, para_idx=para_idx)

    formula_part = match.group(1).strip()
    num_val = match.group(2) if match.group(2) else match.group(3)
    num_part = f"({num_val})"
    if not formula_part.startswith('$'):
        formula_part = f"${formula_part}$"

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(2.2), Cm(10.2), Cm(2.2)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
        table.rows[0].cells[i].width = w

    tblPr = table._tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # 中间单元格：公式居中
    cell_mid = table.rows[0].cells[1]
    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_format(p_mid, line_spacing_pt=20, first_line_indent=Cm(0))
    if not formula_part.startswith('$'):
        formula_part = f"${formula_part}$"
    p_mid.add_run(formula_part)
    if _HAS_OMML:
        replace_latex_with_omml(p_mid, global_stats=global_stats, para_idx=para_idx)

    # 右侧单元格：编号右对齐
    cell_right = table.rows[0].cells[2]
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_format(p_right, line_spacing_pt=20, first_line_indent=Cm(0))
    num_font = getattr(config, "number_font_name", "Times New Roman") if config else "Times New Roman"
    num_size = getattr(config, "number_font_size_pt", 10.5) if config else 10.5
    run_num = p_right.add_run(num_part)
    _set_run_font(run_num, cn_font='宋体', en_font=num_font, size_pt=num_size)

    return p_mid
