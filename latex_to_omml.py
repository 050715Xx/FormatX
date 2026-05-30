#!/usr/bin/env python3
"""
LaTeX 数学公式 → Word OMML 公式转换器

将 $...$ / $$...$$ 包裹的 LaTeX 公式转为 Word 原生公式对象（OMML），
渲染效果与数学编辑器一致：根号带横线、分数有分数线、积分有上下限等。

用法：
    from latex_to_omml import latex_in_text_to_omml
    para = doc.add_paragraph()
    latex_in_text_to_omml(para, '计算 $\\int_0^1 x dx$ 的值')

依赖：pip install python-docx lxml
"""

import re
import uuid
from lxml import etree
from docx.oxml.ns import qn

from core.repair import repair_formula_text
from core.formula_stats import FormulaOccurrence

# OMML 命名空间
M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

# ── LaTeX → Token 流 ──────────────────────────────────────────

# 动态组装命令集：从统一语义注册表加载，按长度降序防截胡
from core.semantics import (
    FUNCTION_NAMES, BIG_OPERATOR_NAMES, GREEK_COMMAND_NAMES,
    BOLD_ITALIC_MARKER_COMMANDS, BOLD_ROMAN_MARKER_COMMANDS,
    ROMAN_MARKER_COMMANDS, ITALIC_MARKER_COMMANDS,
    UPRIGHT_FAMILY_MARKER_COMMANDS,
    RELATION_AND_OPERATOR_SYMBOLS, ARROW_SYMBOLS,
    MISC_MATH_SYMBOLS, FORMATX_LEGACY_COMPAT_COMMANDS,
    STRUCTURAL_COMMANDS,
    DOT_SYMBOLS, EXTRA_ACCENT_COMMANDS,
    PHANTOM_AND_STYLE_COMMANDS, STACKED_AND_EXTENSIBLE_COMMANDS,
    STRUCTURE_DELIMITERS,
)
from core.symbols import LATEX_COMMAND_TO_UNICODE

_ALL_KNOWN_COMMANDS = (
    FUNCTION_NAMES | BIG_OPERATOR_NAMES | GREEK_COMMAND_NAMES |
    BOLD_ITALIC_MARKER_COMMANDS | BOLD_ROMAN_MARKER_COMMANDS |
    ROMAN_MARKER_COMMANDS | ITALIC_MARKER_COMMANDS |
    UPRIGHT_FAMILY_MARKER_COMMANDS |
    RELATION_AND_OPERATOR_SYMBOLS | ARROW_SYMBOLS |
    MISC_MATH_SYMBOLS | FORMATX_LEGACY_COMPAT_COMMANDS |
    STRUCTURAL_COMMANDS |
    DOT_SYMBOLS | EXTRA_ACCENT_COMMANDS |
    PHANTOM_AND_STYLE_COMMANDS | STACKED_AND_EXTENSIBLE_COMMANDS |
    STRUCTURE_DELIMITERS |
    set(LATEX_COMMAND_TO_UNICODE.keys())
)

_LATEX_COMMANDS = sorted(_ALL_KNOWN_COMMANDS, key=lambda x: -len(x))

# 拼接正则
_CMD_PATTERN = '|'.join(re.escape(c) for c in _LATEX_COMMANDS)
_TOKEN = re.compile(
    r'\\text\{[^}]*\}'
    r'|\\(?:' + _CMD_PATTERN + r')(?![a-zA-Z])'
    r'|[&|!_^}{)(\[\]$]'
    r'|\s+'
)


def _tokenize(formula):
    """将 LaTeX 公式字符串切分为 token 列表。"""
    tokens = []
    pos = 0
    while pos < len(formula):
        m = _TOKEN.match(formula, pos)
        if m:
            tok = m.group()
            if tok.startswith('\\text{'):
                tokens.append(tok)
            elif tok.strip():
                tokens.append(tok.strip())
            pos = m.end()
        else:
            tokens.append(formula[pos])
            pos += 1
    return tokens


# ── Token → OMML ──────────────────────────────────────────────

def _omml_element(tag):
    return etree.Element(qn('m:' + tag), nsmap={'m': M})


def _omml_run(text, is_func=False, is_text=False):
    """
    创建 <m:r><m:rPr>...</m:rPr><m:t>text</m:t></m:r>
    is_func: 函数名（sin, cos）→ upright
    is_text: 普通文本（\text{...}）→ 继承段落的宋体小四
    数学变量不加 nor（保持 math italic 默认字体）
    """
    r = _omml_element('r')
    rPr = _omml_element('rPr')
    if is_text:
        nor = _omml_element('nor')
        rPr.append(nor)
    if is_func:
        sty = _omml_element('sty')
        sty.set(qn('m:val'), 'f')
        rPr.append(sty)
    r.append(rPr)
    t = _omml_element('t')
    t.text = text
    r.append(t)
    return r


def _omml_text(text):
    """\text{...} 文本节点，继承段落字体。"""
    return _omml_run(text, is_text=True)
    return r

# ── 函数/符号映射 ──────────────────────────────────────────────

_SYMBOLS = {
    '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
    '\\epsilon': 'ε', '\\varepsilon': 'ε', '\\zeta': 'ζ', '\\eta': 'η',
    '\\theta': 'θ', '\\vartheta': 'ϑ', '\\iota': 'ι', '\\kappa': 'κ',
    '\\lambda': 'λ', '\\mu': 'µ', '\\nu': 'ν', '\\xi': 'ξ',
    '\\pi': 'π', '\\rho': 'ρ', '\\sigma': 'σ', '\\tau': 'τ', '\\upsilon': 'υ',
    '\\phi': 'φ', '\\varphi': 'φ', '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω',
    '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
    '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Upsilon': 'ϒ',
    '\\Phi': 'Φ', '\\Psi': 'Ψ', '\\Omega': 'Ω',
    '\\times': '×', '\\cdot': '·', '\\div': '÷', '\\pm': '±', '\\mp': '∓',
    '\\infty': '∞', '\\partial': '∂', '\\nabla': '∇',
    '\\approx': '≈', '\\equiv': '≡', '\\neq': '≠', '\\leq': '≤', '\\geq': '≥',
    '\\ll': '≪', '\\gg': '≫', '\\propto': '∝', '\\sim': '∼', '\\simeq': '≃',
    '\\parallel': '∥', '\\perp': '⊥', '\\to': '→', '\\rightarrow': '→',
    '\\Rightarrow': '⇒', '\\implies': '⇒', '\\iff': '⇔',
    '\\leftarrow': '←', '\\Leftarrow': '⇐',
    '\\ldots': '…', '\\cdots': '…', '\\vdots': '⋮', '\\ddots': '⋱',
    '\\forall': '∀', '\\exists': '∃', '\\neg': '¬',
    '\\in': '∈', '\\notin': '∉', '\\ni': '∋',
    '\\subset': '⊂', '\\subseteq': '⊆', '\\supset': '⊃', '\\supseteq': '⊇',
    '\\cup': '∪', '\\cap': '∩', '\\emptyset': '∅',
    '\\land': '∧', '\\lor': '∨',
    '\\mid': '|',
    '\\prime': "'", "\\'": "'",
    '\\oplus': '⊕', '\\otimes': '⊗', '\\odot': '⊙',
    '\\ominus': '⊖', '\\oslash': '⊘', '\\star': '⋆',
    '\\circ': '∘', '\\bullet': '•', '\\diamond': '⋄',
    '\\bigoplus': '⨁', '\\bigotimes': '⨂', '\\bigodot': '⨀',
}

_FUNC_NAMES = {
    '\\sin': 'sin', '\\cos': 'cos', '\\tan': 'tan',
    '\\csc': 'csc', '\\sec': 'sec', '\\cot': 'cot',
    '\\arcsin': 'arcsin', '\\arccos': 'arccos', '\\arctan': 'arctan',
    '\\log': 'log', '\\ln': 'ln', '\\exp': 'exp',
    '\\det': 'det', '\\gcd': 'gcd',
    '\\max': 'max', '\\min': 'min', '\\sup': 'sup', '\\inf': 'inf',
    '\\lim': 'lim', '\\limsup': 'limsup', '\\liminf': 'liminf',
}


def _leftright_delim(delim):
    """\left( → ( / \left. → '' 等，转成 OMML 分隔符字符。"""
    dmap = {'(': '(', ')': ')', '[': '[', ']': ']',
            '{': '{', '}': '}', '|': '|', '.': ''}
    return dmap.get(delim, delim)


# ── 递归下降解析器 ────────────────────────────────────────────

class LatexParser:
    def __init__(self, tokens):
        self.tokens = tokens
        self.pos = 0

    def peek(self):
        return self.tokens[self.pos] if self.pos < len(self.tokens) else None

    def consume(self):
        t = self.peek()
        if t is not None:
            self.pos += 1
        return t

    def expect(self, tok):
        t = self.consume()
        if t != tok:
            raise ValueError(f'Expected {tok}, got {t}')

    def parse(self):
        """解析整个公式，返回 OMML 元素列表。"""
        elements = []
        while self.pos < len(self.tokens):
            el = self._parse_expression()
            if el is not None:
                if isinstance(el, list):
                    elements.extend(el)
                else:
                    elements.append(el)
        return elements

    def _parse_expression(self):
        """解析一个表达式（基 + 可选上下标）。"""
        if self.pos >= len(self.tokens):
            return None
        # 解析基础表达式
        base = self._parse_atom()
        if base is None:
            return None
        # 检查后续 ^ 和 _
        has_sub = False
        has_sup = False
        sub_elem = None
        sup_elem = None

        while self.peek() in ('^', '_'):
            kind = self.consume()
            content = self._parse_script_content()
            if kind == '^':
                has_sup = True
                sup_elem = content
            else:
                has_sub = True
                sub_elem = content

        if not has_sub and not has_sup:
            return base

        # 用 sSup / sSub / sSubSup 包裹
        return self._wrap_script(base, sub_elem, sup_elem)

    def _parse_script_content(self):
        """解析 ^ 或 _ 后面的内容。"""
        if self.peek() == '{':
            self.consume()
            inner = self._read_until_brace()
            self.consume()
            el = _omml_element('r')
            for child in LatexParser(inner).parse():
                if isinstance(child, etree._Element):
                    el.append(child)
            return el
        else:
            # 单个 token
            tok = self.consume()
            if tok.startswith('\\'):
                cmd_result = self._process_command(tok)
                e = _omml_element('r')
                if isinstance(cmd_result, etree._Element):
                    # If it's already an m:r element, use as-is
                    if cmd_result.tag == qn('m:r'):
                        return cmd_result
                    e.append(cmd_result)
                elif cmd_result:
                    e.append(_omml_run(str(cmd_result)))
                return e
            return _omml_run(tok)

    def _wrap_script(self, base, sub, sup):
        """用 m:sSup / m:sSub / m:sSubSup 包裹 base + scripts。"""
        if sub is not None and sup is not None:
            elem = _omml_element('sSubSup')
        elif sup is not None:
            elem = _omml_element('sSup')
        else:
            elem = _omml_element('sSub')

        # Base
        e = _omml_element('e')
        if isinstance(base, etree._Element):
            e.append(base)
        else:
            e.append(_omml_run(str(base)))
        elem.append(e)

        # Sub
        if sub is not None:
            s = _omml_element('sub')
            if isinstance(sub, etree._Element):
                s.append(sub)
            elem.append(s)

        # Sup
        if sup is not None:
            s = _omml_element('sup')
            if isinstance(sup, etree._Element):
                s.append(sup)
            elem.append(s)

        return elem

    def _parse_atom(self):
        """解析一个原子表达式（花括号组 / 命令 / 隐式括号 / 普通字符）。"""
        t = self.peek()
        if t is None:
            return None

        if t == '{':
            return self._parse_group()
        if t == '}':
            self.consume()
            return None
        if t == '^' or t == '_':
            return None  # 由上层 _parse_expression 处理
        if t == '&':
            self.consume()
            return None
        # 隐式括号：(...) / [...] → 视为分组，使 ^2 挂在整组上
        if t in ('(', '['):
            delim = t
            close = ')' if t == '(' else ']'
            return self._parse_implicit_group(delim, close)
        if t.startswith('\\text{'):
            self.consume()
            return _omml_text(t[6:-1])
        if t.startswith('\\'):
            self.consume()
            return self._process_command(t)
        # 普通字符
        self.consume()
        return _omml_run(t)

    def _parse_implicit_group(self, open_delim, close_delim):
        """解析 (...) 或 [...] 隐式分组，使 ^2 能挂在整组上。"""
        self.consume()  # open_delim
        depth = 1
        inner_tokens = []
        while self.pos < len(self.tokens) and depth > 0:
            t = self.peek()
            if t == open_delim:
                depth += 1
            elif t == close_delim:
                depth -= 1
                if depth == 0:
                    break
            if depth > 0:
                inner_tokens.append(t)
            self.pos += 1
        self.consume()  # close_delim

        # 用 OMML <m:d> 包裹（和 \left...\right 一样结构）
        d = _omml_element('d')
        dPr = _omml_element('dPr')
        for side, ch in [('begChr', open_delim), ('endChr', close_delim)]:
            el = _omml_element(side)
            el.set(qn('m:val'), ch)
            dPr.append(el)
        d.append(dPr)
        e = _omml_element('e')
        for el in LatexParser(inner_tokens).parse():
            if isinstance(el, etree._Element):
                e.append(el)
        d.append(e)
        return d

    def _parse_group(self):
        self.consume()  # {
        inner = self._read_until_brace()
        self.consume()  # }
        elements = LatexParser(inner).parse()
        if not elements:
            return _omml_run('')
        if len(elements) == 1:
            return elements[0]
        # 返回一个包含所有子元素的 m:r
        grp = _omml_element('r')
        for el in elements:
            if isinstance(el, etree._Element):
                grp.append(el)
        return grp

    def _read_until_brace(self):
        result = []
        depth = 1
        while self.pos < len(self.tokens) and depth > 0:
            t = self.peek()
            if t == '{':
                depth += 1
            elif t == '}':
                depth -= 1
                if depth == 0:
                    break
            if depth > 0:
                result.append(t)
            self.pos += 1
        return result

    def _process_command(self, cmd):
        """处理单个 LaTeX 命令（不含 ^ _ 等）。"""
        if cmd == '\\sqrt':
            return self._parse_sqrt()
        if cmd == '\\frac':
            return self._parse_frac()
        if cmd in ('\\int', '\\sum', '\\prod'):
            return self._parse_nary(cmd)
        # \bar{A} / \hat{A} / \tilde{A} / \vec{A} / \dot{A}
        _ACCENT_CHARS = {
            '\\bar': '̅', '\\hat': '̂', '\\tilde': '̃',
            '\\vec': '⃗', '\\dot': '̇', '\\ddot': '̈',
        }
        if cmd in _ACCENT_CHARS:
            return self._parse_accent(cmd, _ACCENT_CHARS[cmd])

        if cmd in _FUNC_NAMES:
            return _omml_run(_FUNC_NAMES[cmd], is_func=True)
        if cmd == '\\left':
            return self._parse_leftright()
        if cmd == '\\right':
            return None
        if cmd == '\\text':
            if self.peek() == '{':
                self.consume()
                inner = self._read_until_brace()
                self.consume()
                return _omml_text(''.join(inner))
            return _omml_text('')
        if cmd in ('\\boldsymbol', '\\mathbf'):
            base = self._parse_atom()
            if base is None:
                return _omml_run('')
            rpr = base.find(qn('m:rPr'))
            if rpr is None:
                rpr = _omml_element('rPr')
                base.insert(0, rpr)
            sty = _omml_element('sty')
            sty.set(qn('m:val'), 'bi' if cmd == '\\boldsymbol' else 'b')
            rpr.append(sty)
            return base
        if cmd in ('\\,', '\\;', '\\ ', '\\!', '\\quad', '\\qquad'):
            return _omml_run(' ')
        if cmd in _SYMBOLS:
            return _omml_run(_SYMBOLS[cmd])
        return _omml_run(cmd.replace('\\', ''))

    def _parse_sqrt(self):
        self.expect('{')
        inner_tokens = self._read_until_brace()
        self.consume()
        rad = _omml_element('rad')
        deg = _omml_element('deg')
        rad.append(deg)
        e = _omml_element('e')
        for el in LatexParser(inner_tokens).parse():
            if isinstance(el, etree._Element):
                e.append(el)
        rad.append(e)
        return rad

    def _parse_leftright(self):
        """解析 \left( ... \right) 为隐式分组，使 ^2 等脚本挂在整组上。"""
        left_delim = self.consume() if self.peek() else '.'
        # 收集内容直到匹配的 \right
        depth = 1
        inner_tokens = []
        while self.pos < len(self.tokens) and depth > 0:
            t = self.peek()
            if t == '\\left':
                self.consume()
                inner_tokens.append(t)
                inner_tokens.append(self.consume() or '.')
                depth += 1
            elif t == '\\right':
                self.consume()
                if depth == 1:
                    break
                inner_tokens.append(t)
                inner_tokens.append(self.consume() or '.')
                depth -= 1
            else:
                inner_tokens.append(t)
                self.pos += 1
        right_delim = self.consume() if self.pos < len(self.tokens) and self.peek() else '.'
        # 解析内部内容
        inner = LatexParser(inner_tokens).parse()
        # 构建 OMML <m:d> 分隔符元素
        d = _omml_element('d')
        dPr = _omml_element('dPr')
        beg = _omml_element('begChr')
        beg.set(qn('m:val'), _leftright_delim(left_delim))
        dPr.append(beg)
        end = _omml_element('endChr')
        end.set(qn('m:val'), _leftright_delim(right_delim))
        dPr.append(end)
        d.append(dPr)
        e = _omml_element('e')
        for el in inner:
            if isinstance(el, etree._Element):
                e.append(el)
        d.append(e)
        return d

    def _parse_accent(self, cmd, accent_char):
        """解析 \bar{A} / \hat{x} 等重音命令。"""
        if self.peek() == '{':
            self.consume()
            inner = self._read_until_brace()
            self.consume()
            base_el = LatexParser(inner).parse()
            if len(base_el) == 1:
                base = base_el[0]
            else:
                # 多字符底（如 \bar{AB}），包在一个 m:r 里
                base = _omml_element('r')
                for el in base_el:
                    if isinstance(el, etree._Element):
                        base.append(el)
        else:
            base = self._parse_atom()
            if base is None:
                base = _omml_run(accent_char)

        base = base or _omml_run('')
        acc = _omml_element('acc')
        accPr = _omml_element('accPr')
        chr_e = _omml_element('chr')
        chr_e.set(qn('m:val'), accent_char)
        accPr.append(chr_e)
        acc.append(accPr)
        e = _omml_element('e')
        if isinstance(base, etree._Element):
            e.append(base)
        acc.append(e)
        return acc

    def _parse_frac(self):
        self.expect('{')
        num_tokens = self._read_until_brace()
        self.consume()
        self.expect('{')
        den_tokens = self._read_until_brace()
        self.consume()
        f = _omml_element('f')
        fPr = _omml_element('fPr')
        fType = _omml_element('type')
        fType.set(qn('m:val'), 'bar')
        fPr.append(fType)
        f.append(fPr)
        num = _omml_element('num')
        for el in LatexParser(num_tokens).parse():
            if isinstance(el, etree._Element):
                num.append(el)
        f.append(num)
        den = _omml_element('den')
        for el in LatexParser(den_tokens).parse():
            if isinstance(el, etree._Element):
                den.append(el)
        f.append(den)
        return f

    def _parse_nary(self, cmd):
        """解析 \int_a^b 或 \sum_{a}^{b}"""
        nary = _omml_element('nary')
        naryPr = _omml_element('naryPr')
        chr_elem = _omml_element('chr')
        chr_elem.set(qn('m:val'), {
            '\\int': '∫', '\\sum': '∑', '\\prod': '∏'
        }[cmd])
        naryPr.append(chr_elem)
        # 设置 limits 属性
        limLoc = _omml_element('limLoc')
        limLoc.set(qn('m:val'), 'subSup')
        naryPr.append(limLoc)
        nary.append(naryPr)

        sub = None
        sup = None
        if self.peek() == '_':
            self.consume()
            sub = self._parse_script_content()
        if self.peek() == '^':
            self.consume()
            sup = self._parse_script_content()

        if sub is not None:
            s = _omml_element('sub')
            if isinstance(sub, etree._Element):
                s.append(sub)
            nary.append(s)
        if sup is not None:
            s = _omml_element('sup')
            if isinstance(sup, etree._Element):
                s.append(sup)
            nary.append(s)

        # 被积函数
        e = _omml_element('e')
        rest = LatexParser(self.tokens[self.pos:]).parse()
        for el in rest:
            if isinstance(el, etree._Element):
                e.append(el)
        nary.append(e)
        self.pos = len(self.tokens)
        return nary


# ── 公开接口 ──────────────────────────────────────────────────

def _wrap_in_oMath(elements):
    """将解析出的元素列表组装为 Word 原生 m:oMath 对象"""
    omath = etree.Element(qn('m:oMath'), nsmap={'m': M})
    for el in elements:
        if isinstance(el, etree._Element):
            omath.append(el)
    return omath


def _pre_wrap_english_words(formula):
    """拦截预处理：将孤立的英文单词包裹为 \\text{}，防止被 OMML 解析为斜体变量乘积。
    已含 \\text{} 的公式直接原样返回，信任用户的排版意图。"""
    if '\\text{' in formula:
        return formula
    prev = None
    curr = formula
    while prev != curr:
        prev = curr
        curr = re.sub(r'(^|\s)([a-zA-Z]{2,})(?=\s|$)', r'\1\\text{\2}', curr)
    curr = re.sub(r'(\\text\{[a-zA-Z]{2,}\})\s+', r'\1\\ ', curr)
    curr = re.sub(r'\s+(\\text\{[a-zA-Z]{2,}\})', r'\\ \1', curr)
    return curr


def latex_to_omml_element(formula, para_idx=-1, source_hint="plain_text",
                         global_stats=None):
    """将单个 LaTeX 公式字符串转为 OMML XML 元素。

    采用"快慢车道"架构：
    - 快车道：原文直接用 LatexParser 解析，成功即返回（覆盖 98% 正常公式）
    - 慢车道：解析失败时调用本地抢救引擎修复后再解析
    - 兜底：修复无效时原样退回为纯文本 run，绝不崩溃
    """
    occ = FormulaOccurrence(
        formula_id=uuid.uuid4().hex[:6],
        original_text=formula,
        repaired_text=formula,
        source_type=source_hint,
        paragraph_index=para_idx,
        confidence=1.0
    )

    # ── 快车道：先试原文 ──
    try:
        tokens = _tokenize(_pre_wrap_english_words(formula))
        elements = LatexParser(tokens).parse()
        if elements:
            result_omml = _wrap_in_oMath(elements)
            if global_stats is not None:
                global_stats.converted += 1
                global_stats.record_occurrence(occ)
            return result_omml
    except Exception as e:
        occ.warnings.append(f"原公式语法解析失败: {str(e)}")

    # ── 慢车道：启动本地抢救引擎 ──
    try:
        outcome = repair_formula_text(formula)
        if outcome.confidence >= 0.6 and outcome.text != formula:
            tokens = _tokenize(_pre_wrap_english_words(outcome.text))
            elements = LatexParser(tokens).parse()
            if elements:
                result_omml = _wrap_in_oMath(elements)
                occ.repaired_text = outcome.text
                occ.confidence = outcome.confidence
                occ.is_fixed = True
                occ.warnings.extend(outcome.warnings)
                if global_stats is not None:
                    global_stats.converted += 1
                    global_stats.record_occurrence(occ)
                return result_omml
    except Exception as e:
        occ.warnings.append(f"ICU 抢救尝试失败: {str(e)}")

    # ── 兜底：救不活就原样退回为纯文本 run ──
    occ.confidence = 0.35
    occ.warnings.append("公式彻底损坏，已降级为普通纯文本展示")
    if global_stats is not None:
        global_stats.record_occurrence(occ)
    return _omml_run(formula)


def latex_in_text_to_omml(paragraph, text, para_idx=-1, global_stats=None):
    """
    将 text 中的 $...$（行内）和 $$...$$（块级）公式转为 OMML，
    插入到给定 paragraph 中。非公式部分保留为普通 run。

    支持带身世前缀的公式： [OLE_EQUATION]$...$ / [OCR_FRAGMENT]$...$
    """
    split_pattern = (
        r'(\[OLE_EQUATION\]\$\$?[^$]+\$\$?|'
        r'\[OCR_FRAGMENT\]\$\$?[^$]+\$\$?|'
        r'\$\$?[^$]+\$\$?)'
    )
    parts = re.split(split_pattern, text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue

        source_hint = "plain_text"
        if part.startswith("[OLE_EQUATION]"):
            source_hint = "ole_equation"
            part = part.replace("[OLE_EQUATION]", "", 1)
        elif part.startswith("[OCR_FRAGMENT]"):
            source_hint = "ocr_fragment"
            part = part.replace("[OCR_FRAGMENT]", "", 1)

        if part.startswith('$'):
            is_block = part.startswith('$$')
            raw_formula = part.strip('$')

            if raw_formula:
                omml = latex_to_omml_element(
                    raw_formula,
                    para_idx=para_idx,
                    source_hint=source_hint,
                    global_stats=global_stats
                )
                if is_block:
                    run = paragraph.add_run()
                    run._r.append(omml)
                else:
                    run = paragraph.add_run()
                    run._r.append(omml)
        elif part:
            paragraph.add_run(part)


# ── 自动包裹安全门控 ──────────────────────────────────────────────

_LATEX_AUTO_WRAP_CMDS = [
    '\\frac', '\\int', '\\sum', '\\lim', '\\sqrt',
    '\\Delta', '\\alpha', '\\beta', '\\gamma', '\\pi',
]

_HAS_CJK = re.compile(r'[一-龥]')

_ENGLISH_WORDS = re.compile(r'(?<!\\)\b[a-zA-Z]{3,}\b')


def _is_safe_for_auto_wrap(text: str) -> bool:
    """三级门控：仅当段落不含中日韩文字、且纯英文词 < 3 个时，才允许整段 $ 包裹"""
    has_latex = any(cmd in text for cmd in _LATEX_AUTO_WRAP_CMDS)
    if not has_latex:
        return False
    if _HAS_CJK.search(text):
        return False
    if len(_ENGLISH_WORDS.findall(text)) >= 3:
        return False
    return True


_LATEX_FRAGMENT = re.compile(r'(\\[a-zA-Z]+(?:\{[^{}]*\}|\([^(]*\))*)')


def _wrap_latex_fragments(text: str) -> str | None:
    """识别文本中的 LaTeX 片段，逐个包裹为 $...$，合并相邻块后返回。
    找不到任何 LaTeX 特征时返回 None。"""
    found = False

    def _replacer(m):
        nonlocal found
        found = True
        return f"${m.group(1)}$"

    wrapped = _LATEX_FRAGMENT.sub(_replacer, text)
    if not found:
        return None
    # 分离相邻公式边界：$\ln$$\left(...)$ → $\ln$ $\left(...)$
    wrapped = re.sub(r'(?<!\$)(\$[^$]+\$)(\$(?!\$))', r'\1 \2', wrapped)
    return wrapped


def replace_latex_with_omml(paragraph, global_stats=None, para_idx=-1):
    """
    遍历 paragraph 的所有 run，将其中 LaTeX 公式替换为 OMML。
    含公式的段落行距改为"至少 20 磅"，防止高公式与上下行重叠。
    global_stats: 可选 FormulaRuleStats 实例，用于记录公式置信度
    para_idx: 段落索引，用于公式溯源追踪
    """
    full_text = paragraph.text.replace('＄', '$')

    # 在底层直接接管 AI 的 LaTeX 括号，支持跨行公式
    full_text = re.sub(r'\\\[(.*?)\\\]', r'$$\1$$', full_text, flags=re.DOTALL)
    full_text = re.sub(r'\\\((.*?)\\\)', r'$\1$', full_text, flags=re.DOTALL)

    # 清洗 AI 填空题干扰字符 \____ → " . "
    full_text = re.sub(r'\\_+', ' . ', full_text)

    if '$' not in full_text:
        if _is_safe_for_auto_wrap(full_text):
            full_text = f"${full_text.strip()}$"
        else:
            wrapped = _wrap_latex_fragments(full_text)
            if wrapped:
                full_text = wrapped
            else:
                return False

    backup_text = full_text
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    try:
        latex_in_text_to_omml(paragraph, full_text,
                              para_idx=para_idx, global_stats=global_stats)
        if not paragraph.runs and full_text.strip():
            raise ValueError("公式解析器返回了空结果")
    except Exception:
        paragraph.add_run(backup_text)

    # 公式段落：行距改为"至少 20pt"
    pf = paragraph.paragraph_format
    from docx.shared import Pt
    pf.line_spacing = Pt(20)
    from docx.enum.text import WD_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    return True
